from __future__ import annotations

import io
import re
from pathlib import Path

import pdfplumber
from docx import Document

from backend.api.schemas.domain import ResumeDocument, ResumeMetrics, ResumeSection
from backend.services.shared.text_normalization import repair_mojibake


SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "summary": ("summary", "profile", "about", "个人总结", "自我评价", "个人优势", "求职意向"),
    "education": ("education", "教育背景", "教育经历", "学历背景"),
    "experience": (
        "experience",
        "work experience",
        "employment",
        "实习经历",
        "工作经历",
        "职业经历",
        "工作及实习经历",
        "实习和工作经历",
        "工作 / 实习经历",
        "工作/实习经历",
        "专业经历",
    ),
    "projects": ("projects", "project experience", "项目经历", "项目经验", "项目实践"),
    "skills": (
        "skills",
        "technical skills",
        "专业技能",
        "技能",
        "核心技能",
        "证书技能",
        "技能证书",
        "证书与技能",
        "工具技能",
    ),
    "certifications": ("certifications", "licenses", "证书", "资格证", "证书情况"),
    "awards": ("awards", "honors", "获奖", "荣誉", "在校经历", "校内经历", "获奖荣誉"),
}

HEADING_STRIP_RE = re.compile(r"[\s:：/\\|·•\-\(\)\[\]【】（）]+")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(?<!\d)(?:\+?86[- ]?)?1[3-9]\d{9}(?!\d)")
BULLET_RE = re.compile(r"^\s*[-*•◦●]")
TOKEN_RE = re.compile(r"[A-Za-z0-9_+#/.%-]+|[\u4e00-\u9fff]{2,}")
NUMBER_RE = re.compile(r"\d+(?:\.\d+)?%?")
DATE_RANGE_RE = re.compile(
    r"(?:19|20)\d{2}(?:[./年-]\d{1,2})?\s*(?:[-~至到]\s*(?:现在|至今|present|current|(?:19|20)\d{2}(?:[./年-]\d{1,2})?))?",
    re.IGNORECASE,
)
ACTION_VERBS = (
    "built",
    "led",
    "drove",
    "launched",
    "optimized",
    "designed",
    "implemented",
    "managed",
    "analyzed",
    "improved",
    "reduced",
    "increased",
    "delivered",
    "created",
    "developed",
    "负责",
    "主导",
    "搭建",
    "设计",
    "优化",
    "推动",
    "协调",
    "分析",
    "管理",
    "交付",
    "提升",
    "降低",
    "完成",
    "落地",
)


def parse_resume_bytes(filename: str, file_bytes: bytes) -> ResumeDocument:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        raw_text, metrics, warnings = _parse_pdf(file_bytes)
        file_type = "pdf"
    elif extension == ".docx":
        raw_text, metrics, warnings = _parse_docx(file_bytes)
        file_type = "docx"
    else:
        raise ValueError(f"Unsupported resume file type: {extension}")

    repaired_text = repair_mojibake(raw_text)
    if repaired_text != raw_text:
        warnings.append("Applied mojibake repair heuristic to extracted resume text.")
        raw_text = repaired_text

    sections = _segment_sections(raw_text)
    metrics.section_count = len(sections)

    return ResumeDocument(
        filename=filename,
        file_type=file_type,
        raw_text=raw_text,
        sections=sections,
        metrics=metrics,
        inferred_name=_infer_name(raw_text),
        inferred_email=_search_first(EMAIL_RE, raw_text),
        inferred_phone=_search_first(PHONE_RE, raw_text),
        parsing_warnings=warnings,
    )


def _parse_pdf(file_bytes: bytes) -> tuple[str, ResumeMetrics, list[str]]:
    text_parts: list[str] = []
    warnings: list[str] = []
    page_count = 0
    table_count = 0
    image_count = 0
    text_pages = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                text_pages += 1
                text_parts.append(text)
            try:
                table_count += len(page.find_tables())
            except Exception:
                warnings.append("Table detection failed on at least one PDF page.")
            image_count += len(page.images or [])

    raw_text = "\n\n".join(text_parts).strip()
    metrics = _build_metrics(
        raw_text=raw_text,
        page_count=page_count,
        table_count=table_count,
        image_count=image_count,
        text_page_ratio=(text_pages / page_count) if page_count else 0.0,
    )
    if not raw_text:
        warnings.append("No text could be extracted from the PDF.")
    return raw_text, metrics, warnings


def _parse_docx(file_bytes: bytes) -> tuple[str, ResumeMetrics, list[str]]:
    warnings: list[str] = []
    document = Document(io.BytesIO(file_bytes))

    paragraph_lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    raw_text = "\n".join(paragraph_lines + table_lines).strip()
    image_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    metrics = _build_metrics(
        raw_text=raw_text,
        page_count=1,
        table_count=len(document.tables),
        image_count=image_count,
        text_page_ratio=1.0 if raw_text else 0.0,
    )
    if not raw_text:
        warnings.append("No text could be extracted from the DOCX file.")
    return raw_text, metrics, warnings


def _build_metrics(
    *,
    raw_text: str,
    page_count: int,
    table_count: int,
    image_count: int,
    text_page_ratio: float,
) -> ResumeMetrics:
    tokens = TOKEN_RE.findall(raw_text)
    bullet_count = sum(1 for line in raw_text.splitlines() if BULLET_RE.match(line))
    numeric_token_count = len(NUMBER_RE.findall(raw_text))
    date_range_count = len(DATE_RANGE_RE.findall(raw_text))
    action_verb_count = _count_action_verbs(raw_text)
    return ResumeMetrics(
        page_count=page_count,
        word_count=len(tokens),
        bullet_count=bullet_count,
        numeric_token_count=numeric_token_count,
        date_range_count=date_range_count,
        action_verb_count=action_verb_count,
        table_count=table_count,
        image_count=image_count,
        text_page_ratio=round(text_page_ratio, 2),
    )


def _segment_sections(raw_text: str) -> list[ResumeSection]:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return []

    sections: list[ResumeSection] = []
    current_heading = "Header"
    current_type = "header"
    buffer: list[str] = []

    def flush() -> None:
        content = "\n".join(buffer).strip()
        if not content:
            return
        sections.append(
            ResumeSection(
                section_type=current_type,
                heading=current_heading,
                content=content,
                bullet_count=sum(1 for line in buffer if BULLET_RE.match(line)),
            )
        )

    for line in lines:
        section_type = _match_heading(line)
        if section_type:
            flush()
            current_heading = line.rstrip(":：")
            current_type = section_type
            buffer = []
            continue
        buffer.append(line)

    flush()
    if len(sections) == 1 and sections[0].section_type == "header":
        sections[0].section_type = "general"
        sections[0].heading = "General"
    return sections


def _match_heading(line: str) -> str | None:
    normalized = HEADING_STRIP_RE.sub("", line).lower()
    if len(normalized) > 20:
        return None
    for section_type, aliases in SECTION_ALIASES.items():
        if any(HEADING_STRIP_RE.sub("", alias).lower() == normalized for alias in aliases):
            return section_type
    return None


def _infer_name(raw_text: str) -> str | None:
    for line in raw_text.splitlines():
        candidate = line.strip()
        if not candidate:
            continue
        if _match_heading(candidate):
            continue
        if EMAIL_RE.search(candidate) or PHONE_RE.search(candidate):
            continue
        if len(candidate) <= 30:
            return candidate
    return None


def _search_first(pattern: re.Pattern[str], text: str) -> str | None:
    match = pattern.search(text)
    return match.group(0) if match else None


def _count_action_verbs(raw_text: str) -> int:
    lowered = raw_text.lower()
    count = 0
    for verb in ACTION_VERBS:
        if re.fullmatch(r"[a-z ]+", verb):
            count += len(re.findall(rf"\b{re.escape(verb)}\b", lowered))
        else:
            count += lowered.count(verb)
    return count
