"""
========================================
简历文件解析模块
========================================
功能:
  1. 从 PDF 文件中提取文字
  2. 从 Word (.docx) 文件中提取文字
  3. 格式体检：检测表格、图片、多栏等 ATS 不友好元素
========================================
"""

import pdfplumber
from docx import Document
import io


def extract_text_from_pdf(uploaded_file):
    """
    从上传的 PDF 文件中提取全部文字内容。
    
    参数:
        uploaded_file: Streamlit 的 UploadedFile 对象（PDF 格式）
    
    返回:
        str: 提取到的文字内容
    """
    text_parts = []  # 用来收集每一页的文字
    
    # pdfplumber 需要一个文件流，Streamlit 上传的文件本身就是流
    with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    # 把所有页的文字用换行连接起来
    full_text = "\n\n".join(text_parts)
    
    # 重置文件指针（以便后续格式体检时还能再次读取）
    uploaded_file.seek(0)
    
    return full_text


def extract_text_from_docx(uploaded_file):
    """
    从上传的 Word (.docx) 文件中提取全部文字内容。
    
    参数:
        uploaded_file: Streamlit 的 UploadedFile 对象（DOCX 格式）
    
    返回:
        str: 提取到的文字内容
    """
    # python-docx 需要一个文件流
    doc = Document(io.BytesIO(uploaded_file.read()))
    
    # 提取所有段落的文字
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
    
    full_text = "\n".join(text_parts)
    
    # 重置文件指针
    uploaded_file.seek(0)
    
    return full_text


def check_format_issues(uploaded_file, file_type):
    """
    格式体检：检测简历中可能导致 ATS 解析失败的格式问题。
    
    现代 ATS 系统极度排斥：
    - 双栏/多栏排版（文字提取顺序会错乱）
    - 表格（信息可能被打散）
    - 图片（完全无法识别文字）
    
    参数:
        uploaded_file: Streamlit 的 UploadedFile 对象
        file_type: 文件类型，"pdf" 或 "docx"
    
    返回:
        list[dict]: 问题列表，每个问题包含 level（严重程度）和 message（描述）
            level: "error"（严重）, "warning"（警告）, "info"（提示）
    """
    issues = []
    
    if file_type == "pdf":
        issues = _check_pdf_format(uploaded_file)
    elif file_type == "docx":
        issues = _check_docx_format(uploaded_file)
    
    return issues


def _check_pdf_format(uploaded_file):
    """
    检查 PDF 简历的格式问题。
    """
    issues = []
    
    with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
        total_pages = len(pdf.pages)
        total_tables = 0
        total_images = 0
        has_text_issues = False
        
        for page_num, page in enumerate(pdf.pages, start=1):
            # ---- 检测表格 ----
            tables = page.find_tables()
            if tables:
                total_tables += len(tables)
            
            # ---- 检测图片 ----
            images = page.images
            if images:
                total_images += len(images)
            
            # ---- 检测文字提取质量（可能是双栏或扫描件） ----
            text = page.extract_text()
            if not text or len(text.strip()) < 20:
                has_text_issues = True
        
        # 【页数检查】简历超过 2 页通常不被推荐
        if total_pages > 2:
            issues.append({
                "level": "warning",
                "message": f"📄 简历共 {total_pages} 页，建议控制在 1~2 页以内。过长的简历会降低 HR 的阅读耐心。"
            })
        
        # 【表格检查】ATS 对表格支持很差
        if total_tables > 0:
            issues.append({
                "level": "error",
                "message": f"📊 检测到 {total_tables} 个表格元素。ATS 系统解析表格时容易将信息打散或错位，强烈建议改为纯文本列表格式。"
            })
        
        # 【图片检查】ATS 完全无法识别图片中的文字
        if total_images > 0:
            issues.append({
                "level": "warning",
                "message": f"🖼️ 检测到 {total_images} 张图片。ATS 无法识别图片中的内容（包括技能图标、评分条、头像等），建议删除或替换为文字。"
            })
        
        # 【文字提取异常】可能是扫描件或双栏排版
        if has_text_issues:
            issues.append({
                "level": "error",
                "message": "⚠️ 部分页面无法正常提取文字，可能原因：\n"
                           "① 简历为扫描件/截图（ATS 完全无法读取）\n"
                           "② 使用了双栏排版（文字提取顺序会混乱）\n"
                           "③ 使用了特殊字体嵌入\n"
                           "👉 建议：转为单栏纯文本格式的 PDF"
            })
    
    # 如果没有任何问题，给一个正面反馈
    if not issues:
        issues.append({
            "level": "info",
            "message": "✅ 格式体检通过！未发现明显的 ATS 不兼容元素，格式基本规范。"
        })
    
    # 重置文件指针
    uploaded_file.seek(0)
    
    return issues


def _check_docx_format(uploaded_file):
    """
    检查 Word (.docx) 简历的格式问题。
    """
    issues = []
    doc = Document(io.BytesIO(uploaded_file.read()))
    
    # ---- 检测表格 ----
    table_count = len(doc.tables)
    if table_count > 0:
        issues.append({
            "level": "error",
            "message": f"📊 检测到 {table_count} 个表格。许多 ATS 系统在解析 Word 表格时会丢失数据或顺序错乱，建议改为纯文本列表格式。"
        })
    
    # ---- 检测图片/嵌入对象 ----
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    
    if image_count > 0:
        issues.append({
            "level": "warning",
            "message": f"🖼️ 检测到 {image_count} 张图片/图形。ATS 无法识别图片内容，建议删除头像、技能图标、评分条等视觉元素。"
        })
    
    # ---- 检测段落数量（粗略判断内容量） ----
    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty_paras) < 5:
        issues.append({
            "level": "warning",
            "message": "📝 简历内容段落较少，可能信息不够充分。确认是否有内容被放在文本框或表格中（这些部分 ATS 可能无法读取）。"
        })
    
    # 如果没有任何问题，给一个正面反馈
    if not issues:
        issues.append({
            "level": "info",
            "message": "✅ 格式体检通过！Word 文档结构清晰，未发现明显的 ATS 不兼容元素。"
        })
    
    # 重置文件指针
    uploaded_file.seek(0)
    
    return issues
